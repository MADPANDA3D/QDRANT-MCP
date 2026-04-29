from __future__ import annotations

import csv
import io
import math
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class DocumentSection:
    text: str
    page_start: int | None = None
    page_end: int | None = None
    section_heading: str | None = None


@dataclass
class ExtractionResult:
    sections: list[DocumentSection]
    page_count: int | None = None
    warnings: list[str] = field(default_factory=list)
    title_hint: str | None = None
    coverage_ratio: float | None = None
    coverage_pages_total: int | None = None
    coverage_pages_meeting_threshold: int | None = None
    ocr_low_text_threshold_chars: int | None = None
    ocr_min_coverage_ratio: float | None = None
    ocr_budget_pages: int | None = None
    ocr_attempted_pages: int = 0
    low_text_pages_before_ocr: int | None = None
    low_text_pages_after_ocr: int | None = None


_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_CHAPTER_HEADING_RE = re.compile(
    r"\bchapter\s+(\d{1,3})(?:\s*[:.\-]\s*([^\n\r]{1,200}))?",
    flags=re.IGNORECASE,
)


def chunk_text_with_overlap(text: str, chunk_size: int, overlap: int) -> list[str]:
    text = text.strip()
    if not text:
        return []
    if chunk_size <= 0 or len(text) <= chunk_size:
        return [text]
    if overlap < 0:
        overlap = 0
    if overlap >= chunk_size:
        overlap = max(0, chunk_size - 1)

    chunks: list[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        split = end
        if end < length:
            space = text.rfind(" ", start, end)
            if space > start:
                split = space
        chunk = text[start:split].strip()
        if chunk:
            chunks.append(chunk)
        if split >= length:
            break
        next_start = split - overlap
        if next_start <= start:
            next_start = split
        start = next_start

    return chunks or [text[:chunk_size]]


def normalize_text_for_chunking(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def normalize_chapter_map(
    chapter_map: list[dict[str, Any]] | None,
) -> list[dict[str, Any]]:
    if not chapter_map:
        return []

    normalized: list[dict[str, Any]] = []
    for idx, item in enumerate(chapter_map):
        if not isinstance(item, dict):
            raise ValueError(f"chapter_map[{idx}] must be an object.")

        start_page = item.get("start_page")
        end_page = item.get("end_page")
        chapter = item.get("chapter")
        chapter_title = item.get("chapter_title")

        if not isinstance(start_page, int) or start_page <= 0:
            raise ValueError(
                f"chapter_map[{idx}].start_page must be a positive integer."
            )
        if end_page is not None and (
            not isinstance(end_page, int) or end_page < start_page
        ):
            raise ValueError(
                f"chapter_map[{idx}].end_page must be >= start_page when provided."
            )
        if chapter is not None and (not isinstance(chapter, int) or chapter <= 0):
            raise ValueError(f"chapter_map[{idx}].chapter must be a positive integer.")
        if chapter_title is not None and not isinstance(chapter_title, str):
            raise ValueError(f"chapter_map[{idx}].chapter_title must be a string.")

        normalized.append(
            {
                "start_page": start_page,
                "end_page": end_page,
                "chapter": chapter,
                "chapter_title": chapter_title.strip() if chapter_title else None,
            }
        )

    normalized.sort(key=lambda entry: entry["start_page"])
    return normalized


def detect_pdf_chapter_markers(sections: list[DocumentSection]) -> list[dict[str, Any]]:
    markers: list[dict[str, Any]] = []
    seen_pages: set[int] = set()
    for section in sections:
        if section.page_start is None or section.page_start in seen_pages:
            continue
        seen_pages.add(section.page_start)
        preview = section.text[:600]
        match = _CHAPTER_HEADING_RE.search(preview)
        if not match:
            continue
        chapter_number = int(match.group(1))
        title = match.group(2).strip() if match.group(2) else None
        markers.append(
            {
                "start_page": section.page_start,
                "chapter": chapter_number,
                "chapter_title": title,
            }
        )

    markers.sort(key=lambda entry: entry["start_page"])
    return markers


def resolve_chapter_metadata_for_page(
    page: int | None,
    *,
    chapter_map: list[dict[str, Any]] | None = None,
    detected_markers: list[dict[str, Any]] | None = None,
) -> tuple[int | None, str | None]:
    if page is None:
        return None, None

    if chapter_map:
        for item in chapter_map:
            start_page = item["start_page"]
            end_page = item.get("end_page")
            if page < start_page:
                continue
            if end_page is not None and page > end_page:
                continue
            return item.get("chapter"), item.get("chapter_title")

    if detected_markers:
        selected: dict[str, Any] | None = None
        for marker in detected_markers:
            if page >= marker["start_page"]:
                selected = marker
            else:
                break
        if selected:
            return selected.get("chapter"), selected.get("chapter_title")

    return None, None


def decode_bytes_to_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        return data.decode("utf-8"), warnings
    except UnicodeDecodeError:
        decoded = data.decode("utf-8", errors="replace")
        if "\ufffd" in decoded:
            warnings.append("Text contained invalid utf-8 bytes; replacements applied.")
        return decoded, warnings


def count_text_chars(text: str) -> int:
    return len((text or "").strip())


def summarize_pdf_text_coverage(
    page_texts: list[str], low_text_threshold_chars: int
) -> dict[str, Any]:
    threshold = max(0, low_text_threshold_chars)
    total_pages = len(page_texts)
    low_text_pages = [
        index
        for index, text in enumerate(page_texts)
        if count_text_chars(text) < threshold
    ]
    good_pages = total_pages - len(low_text_pages)
    coverage_ratio = (good_pages / total_pages) if total_pages > 0 else 0.0
    return {
        "total_pages": total_pages,
        "good_pages": good_pages,
        "low_text_pages": low_text_pages,
        "coverage_ratio": coverage_ratio,
    }


def compute_ocr_page_budget(
    total_pages: int,
    *,
    ocr_max_pages: int,
    ocr_max_page_ratio: float,
) -> int:
    if total_pages <= 0:
        return 0
    max_pages = max(0, ocr_max_pages)
    ratio = max(0.0, ocr_max_page_ratio)
    ratio_budget = math.ceil(total_pages * ratio)
    if ratio > 0 and ratio_budget == 0:
        ratio_budget = 1
    return max(0, min(max_pages, ratio_budget))


def select_ocr_candidate_pages(
    page_texts: list[str],
    *,
    low_text_threshold_chars: int,
    budget_pages: int,
) -> list[int]:
    if budget_pages <= 0:
        return []
    summary = summarize_pdf_text_coverage(page_texts, low_text_threshold_chars)
    low_pages = summary["low_text_pages"]
    ranked_pages = sorted(
        low_pages, key=lambda index: count_text_chars(page_texts[index])
    )
    return ranked_pages[:budget_pages]


def extract_plain_text(text: str) -> ExtractionResult:
    cleaned = text.strip()
    return ExtractionResult(
        sections=[DocumentSection(text=cleaned)] if cleaned else [],
    )


def extract_markdown_sections(text: str) -> ExtractionResult:
    sections: list[DocumentSection] = []
    buffer: list[str] = []
    current_heading: str | None = None
    title_hint: str | None = None

    def flush():
        nonlocal buffer
        content = "\n".join(buffer).strip()
        if content:
            sections.append(
                DocumentSection(text=content, section_heading=current_heading)
            )
        buffer = []

    for line in text.splitlines():
        match = _MARKDOWN_HEADING_RE.match(line.strip())
        if match:
            flush()
            current_heading = match.group(2).strip()
            if title_hint is None:
                title_hint = current_heading or title_hint
            buffer.append(line.strip())
        else:
            buffer.append(line)

    flush()

    if not sections and text.strip():
        sections.append(DocumentSection(text=text.strip(), section_heading=None))

    return ExtractionResult(sections=sections, title_hint=title_hint)


def _extract_docx_sections_sync(data: bytes) -> ExtractionResult:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise RuntimeError("python-docx is required for .docx files.") from exc

    document = docx.Document(io.BytesIO(data))
    sections: list[DocumentSection] = []
    buffer: list[str] = []
    current_heading: str | None = None
    title_hint: str | None = None

    def flush():
        nonlocal buffer
        content = "\n".join(buffer).strip()
        if content:
            sections.append(
                DocumentSection(text=content, section_heading=current_heading)
            )
        buffer = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = paragraph.style.name or ""
        except Exception:  # pragma: no cover - style access varies
            style_name = ""
        if style_name.lower().startswith("heading"):
            flush()
            current_heading = text
            if title_hint is None:
                title_hint = current_heading
            buffer.append(text)
        else:
            buffer.append(text)

    flush()

    if not sections and document.paragraphs:
        full_text = "\n".join(
            paragraph.text for paragraph in document.paragraphs if paragraph.text
        ).strip()
        if full_text:
            sections.append(DocumentSection(text=full_text, section_heading=None))

    return ExtractionResult(sections=sections, title_hint=title_hint)


def _extract_pdf_text_with_pdftotext(
    data: bytes | None = None,
    *,
    pdf_path: str | Path | None = None,
) -> tuple[str, list[str]]:
    warnings: list[str] = []
    temp_pdf_path = None
    pdf_input_path: str | None = None
    txt_path = None
    try:
        if pdf_path is not None:
            pdf_input_path = str(pdf_path)
        else:
            if data is None:
                raise ValueError("PDF extraction requires either data or pdf_path.")
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
                pdf_file.write(data)
                pdf_file.flush()
                temp_pdf_path = pdf_file.name
                pdf_input_path = temp_pdf_path
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as txt_file:
            txt_path = txt_file.name
        result = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", pdf_input_path, txt_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except FileNotFoundError:  # pragma: no cover - binary missing
        warnings.append("pdftotext is required for PDF fallback extraction.")
        return "", warnings
    finally:
        if temp_pdf_path and os.path.exists(temp_pdf_path):
            try:
                os.unlink(temp_pdf_path)
            except OSError:
                warnings.append("Failed to remove temporary PDF file.")

    if result.returncode != 0:
        warnings.append(
            f"pdftotext exited with code {result.returncode}: "
            f"{result.stderr.decode('utf-8', errors='replace').strip()}"
        )

    output_bytes = b""
    if txt_path and os.path.exists(txt_path):
        try:
            with open(txt_path, "rb") as handle:
                output_bytes = handle.read()
        except OSError:
            warnings.append("Failed to read pdftotext output.")
    else:
        warnings.append("pdftotext produced no output.")

    if txt_path and os.path.exists(txt_path):
        try:
            os.unlink(txt_path)
        except OSError:
            warnings.append("Failed to remove temporary pdftotext output.")

    text, decode_warnings = decode_bytes_to_text(output_bytes)
    warnings.extend(decode_warnings)
    return text, warnings


def _extract_pdf_sections_sync(
    data: bytes | None = None,
    *,
    file_path: str | Path | None = None,
    ocr: bool,
    ocr_low_text_threshold_chars: int = 120,
    ocr_min_coverage_ratio: float = 0.85,
    ocr_max_pages: int = 120,
    ocr_max_page_ratio: float = 0.30,
) -> ExtractionResult:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise RuntimeError("pypdf is required for .pdf files.") from exc

    if data is None and file_path is None:
        raise ValueError("PDF extraction requires either data or file_path.")

    page_texts: list[str] = []
    warnings: list[str] = []
    page_count: int | None = None
    pdf_input_path: str | None = str(file_path) if file_path is not None else None

    try:
        if pdf_input_path is not None:
            reader = PdfReader(pdf_input_path, strict=False)
        else:
            reader = PdfReader(io.BytesIO(data or b""), strict=False)
        page_count = len(reader.pages)
    except Exception as exc:  # pragma: no cover - PDF parsing errors vary
        warnings.append(f"pypdf failed to read PDF: {exc}")
        reader = None

    if reader is not None:
        for idx, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as exc:  # pragma: no cover - PDF parsing errors vary
                warnings.append(f"PDF text extraction failed for page {idx + 1}: {exc}")
                text = ""
            page_texts.append(text)

    def _merge_text(primary_text: str, ocr_text: str) -> str:
        primary = primary_text.strip()
        ocr_clean = ocr_text.strip()
        if not ocr_clean:
            return primary
        if not primary:
            return ocr_clean
        if ocr_clean in primary:
            return primary
        if primary in ocr_clean:
            return ocr_clean
        return f"{primary}\n\n{ocr_clean}"

    coverage_target_ratio = max(0.0, min(1.0, ocr_min_coverage_ratio))
    coverage_before = summarize_pdf_text_coverage(
        page_texts, ocr_low_text_threshold_chars
    )
    low_text_pages_before_ocr = len(coverage_before["low_text_pages"])
    low_text_pages_after_ocr = low_text_pages_before_ocr
    ocr_budget_pages = 0
    ocr_attempted_pages = 0

    if ocr:
        try:
            import pytesseract  # type: ignore
            from pdf2image import convert_from_bytes, convert_from_path  # type: ignore
        except ImportError:  # pragma: no cover - optional dependency missing
            warnings.append("OCR requested but pdf2image/pytesseract not installed.")
        else:
            try:
                ocr_budget_pages = compute_ocr_page_budget(
                    len(page_texts),
                    ocr_max_pages=ocr_max_pages,
                    ocr_max_page_ratio=ocr_max_page_ratio,
                )
                if (
                    coverage_before["coverage_ratio"] < coverage_target_ratio
                    and ocr_budget_pages > 0
                ):
                    ocr_pages = select_ocr_candidate_pages(
                        page_texts,
                        low_text_threshold_chars=ocr_low_text_threshold_chars,
                        budget_pages=ocr_budget_pages,
                    )
                    for page_index in ocr_pages:
                        ocr_attempted_pages += 1
                        images: list[Any] = []
                        try:
                            page_number = page_index + 1
                            if pdf_input_path is not None:
                                images = convert_from_path(
                                    pdf_input_path,
                                    first_page=page_number,
                                    last_page=page_number,
                                )
                            else:
                                images = convert_from_bytes(
                                    data or b"",
                                    first_page=page_number,
                                    last_page=page_number,
                                )
                            if not images:
                                warnings.append(
                                    f"OCR conversion produced no image for page {page_number}."
                                )
                                continue
                            ocr_text = pytesseract.image_to_string(images[0])
                            if ocr_text.strip():
                                page_texts[page_index] = _merge_text(
                                    page_texts[page_index], ocr_text
                                )
                            elif not page_texts[page_index].strip():
                                warnings.append(
                                    f"OCR produced no text for page {page_number}."
                                )
                        finally:
                            for image in images:
                                close = getattr(image, "close", None)
                                if callable(close):
                                    close()
                    coverage_after_ocr = summarize_pdf_text_coverage(
                        page_texts, ocr_low_text_threshold_chars
                    )
                    low_text_pages_after_ocr = len(coverage_after_ocr["low_text_pages"])
                    if coverage_after_ocr["coverage_ratio"] < coverage_target_ratio:
                        warnings.append(
                            "OCR coverage target not met within configured OCR budget. "
                            f"coverage={coverage_after_ocr['coverage_ratio']:.3f}, "
                            f"target={coverage_target_ratio:.3f}, "
                            f"budget={ocr_budget_pages}, attempted={ocr_attempted_pages}."
                        )
                elif coverage_before["coverage_ratio"] < coverage_target_ratio:
                    warnings.append(
                        "OCR was enabled but no OCR budget was available under current limits."
                    )
            except Exception as exc:  # pragma: no cover - OCR errors vary
                warnings.append(f"OCR failed: {exc}")

    coverage_after = summarize_pdf_text_coverage(
        page_texts, ocr_low_text_threshold_chars
    )
    low_text_pages_after_ocr = len(coverage_after["low_text_pages"])

    sections: list[DocumentSection] = []
    for idx, text in enumerate(page_texts):
        if not text.strip():
            continue
        page_number = idx + 1
        sections.append(
            DocumentSection(
                text=text.strip(),
                page_start=page_number,
                page_end=page_number,
                section_heading=None,
            )
        )

    if not sections:
        fallback_text, fallback_warnings = _extract_pdf_text_with_pdftotext(
            data, pdf_path=pdf_input_path
        )
        warnings.extend(fallback_warnings)
        cleaned_fallback = fallback_text.strip()
        if cleaned_fallback:
            sections.append(DocumentSection(text=cleaned_fallback))
        else:
            warnings.append("No text extracted from PDF.")

    return ExtractionResult(
        sections=sections,
        page_count=page_count or len(page_texts),
        warnings=warnings,
        coverage_ratio=coverage_after["coverage_ratio"]
        if coverage_after["total_pages"] > 0
        else None,
        coverage_pages_total=coverage_after["total_pages"],
        coverage_pages_meeting_threshold=coverage_after["good_pages"],
        ocr_low_text_threshold_chars=max(0, ocr_low_text_threshold_chars),
        ocr_min_coverage_ratio=coverage_target_ratio,
        ocr_budget_pages=ocr_budget_pages,
        ocr_attempted_pages=ocr_attempted_pages,
        low_text_pages_before_ocr=low_text_pages_before_ocr,
        low_text_pages_after_ocr=low_text_pages_after_ocr,
    )


def _extract_doc_sections_sync(data: bytes) -> ExtractionResult:
    warnings: list[str] = []
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(data)
            temp_file.flush()
            temp_path = temp_file.name
        result = subprocess.run(
            ["antiword", temp_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except FileNotFoundError as exc:  # pragma: no cover - binary missing
        raise RuntimeError("antiword is required for .doc files.") from exc
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                warnings.append("Failed to remove temporary .doc file.")

    if result.returncode != 0:
        warnings.append(
            f"antiword exited with code {result.returncode}: "
            f"{result.stderr.decode('utf-8', errors='replace').strip()}"
        )

    text, decode_warnings = decode_bytes_to_text(result.stdout)
    warnings.extend(decode_warnings)

    sections: list[DocumentSection] = []
    cleaned = text.strip()
    if cleaned:
        sections.append(DocumentSection(text=cleaned))
    else:
        warnings.append("No text extracted from .doc file.")

    return ExtractionResult(sections=sections, warnings=warnings)


def _extract_csv_sections_sync(data: bytes) -> ExtractionResult:
    text, warnings = decode_bytes_to_text(data)
    if text.startswith("\ufeff"):
        text = text.lstrip("\ufeff")
        warnings.append("CSV UTF-8 BOM removed.")

    sample = text[:8192]
    has_header = False
    dialect: csv.Dialect | type[csv.Dialect] = csv.excel
    if sample.strip():
        try:
            sniffed = csv.Sniffer()
            dialect = sniffed.sniff(sample)
            try:
                has_header = sniffed.has_header(sample)
            except csv.Error:
                has_header = False
        except csv.Error:
            warnings.append("CSV delimiter detection failed; defaulted to comma.")

    reader = csv.reader(io.StringIO(text), dialect=dialect)
    rows = [row for row in reader if any(cell.strip() for cell in row)]
    if not rows:
        warnings.append("No non-empty rows found in CSV.")
        return ExtractionResult(sections=[], warnings=warnings)

    sections: list[DocumentSection] = []
    start_idx = 0
    headers: list[str] = []
    if has_header:
        headers = [
            cell.strip() or f"column_{idx + 1}" for idx, cell in enumerate(rows[0])
        ]
        start_idx = 1
        if not headers:
            headers = []
            start_idx = 0

    for row_idx, row in enumerate(rows[start_idx:], start=start_idx + 1):
        cleaned = [cell.strip() for cell in row]
        if headers:
            parts: list[str] = []
            for col_idx, cell in enumerate(cleaned):
                key = (
                    headers[col_idx]
                    if col_idx < len(headers)
                    else f"column_{col_idx + 1}"
                )
                parts.append(f"{key}: {cell}")
            row_text = " | ".join(parts)
        else:
            row_text = " | ".join(cleaned)
        row_text = row_text.strip().strip("|").strip()
        if not row_text:
            continue
        sections.append(
            DocumentSection(
                text=f"CSV row {row_idx}: {row_text}",
                section_heading="csv",
            )
        )

    if not sections:
        warnings.append("No text rows extracted from CSV.")

    return ExtractionResult(
        sections=sections,
        page_count=len(rows),
        warnings=warnings,
        title_hint=headers[0] if headers else None,
    )


def extract_document_sections(
    file_type: str,
    *,
    text: str | None = None,
    data: bytes | None = None,
    file_path: str | Path | None = None,
    ocr: bool = False,
    ocr_low_text_threshold_chars: int = 120,
    ocr_min_coverage_ratio: float = 0.85,
    ocr_max_pages: int = 120,
    ocr_max_page_ratio: float = 0.30,
) -> ExtractionResult:
    if file_type in {"txt", "text"}:
        if text is None and data is not None:
            text, warnings = decode_bytes_to_text(data)
            result = extract_plain_text(text)
            result.warnings.extend(warnings)
            return result
        return extract_plain_text(text or "")

    if file_type in {"md", "markdown"}:
        if text is None and data is not None:
            text, warnings = decode_bytes_to_text(data)
            result = extract_markdown_sections(text)
            result.warnings.extend(warnings)
            return result
        return extract_markdown_sections(text or "")

    if file_type == "csv":
        if data is None and text is not None:
            data = text.encode("utf-8")
        if data is None:
            raise ValueError(".csv ingestion requires text or binary data.")
        return _extract_csv_sections_sync(data)

    if file_type == "docx":
        if data is None:
            raise ValueError(".docx ingestion requires binary data.")
        return _extract_docx_sections_sync(data)

    if file_type == "doc":
        if data is None:
            raise ValueError(".doc ingestion requires binary data.")
        return _extract_doc_sections_sync(data)

    if file_type == "pdf":
        if data is None and file_path is None:
            raise ValueError(".pdf ingestion requires binary data or file_path.")
        return _extract_pdf_sections_sync(
            data,
            file_path=file_path,
            ocr=ocr,
            ocr_low_text_threshold_chars=ocr_low_text_threshold_chars,
            ocr_min_coverage_ratio=ocr_min_coverage_ratio,
            ocr_max_pages=ocr_max_pages,
            ocr_max_page_ratio=ocr_max_page_ratio,
        )

    raise ValueError(f"Unsupported file_type: {file_type}")
